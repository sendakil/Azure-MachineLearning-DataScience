import docx
import sys
from docx import *

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


file01= "c:\\Me\\Python\\MySample01.docx"

text1 =getText(file01)
doc1 = Document(file01)

tables = doc1.tables

for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)

print(text1)


inline_shapes = doc1.body.InlineShapes
for inline_shape in inline_shapes:
    print(inline_shape.type)





